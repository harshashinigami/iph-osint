"""
E-SitRep Report Generator — PDF and DOCX output.
Generates intelligence situation reports from platform data.
"""
from datetime import datetime, timedelta
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func
from app.models import RawPost, ProcessedPost, Entity, Alert
import os

REPORTS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "reports_output")
os.makedirs(REPORTS_DIR, exist_ok=True)


async def gather_report_data(db: AsyncSession, days: int = 7) -> dict:
    """Gather all data needed for an E-SitRep."""
    since = datetime.utcnow() - timedelta(days=days)

    total_posts = (await db.execute(select(func.count(RawPost.id)).where(RawPost.collected_at >= since))).scalar() or 0
    total_entities = (await db.execute(select(func.count(Entity.id)))).scalar() or 0

    # Alerts by severity
    alert_rows = await db.execute(
        select(Alert.severity, func.count(Alert.id).label("count"))
        .where(Alert.created_at >= since)
        .group_by(Alert.severity)
    )
    alerts_by_severity = {r.severity: r.count for r in alert_rows.fetchall()}

    # Top threat entities
    top_entities = await db.execute(
        select(Entity).order_by(Entity.risk_score.desc()).limit(10)
    )
    entities = [
        {"type": e.entity_type, "value": e.value, "risk": e.risk_score, "mentions": e.mention_count}
        for e in top_entities.scalars().all()
    ]

    # Sentiment breakdown
    sent_rows = await db.execute(
        select(ProcessedPost.sentiment_label, func.count(ProcessedPost.id).label("count"))
        .where(ProcessedPost.sentiment_label.isnot(None))
        .group_by(ProcessedPost.sentiment_label)
    )
    sentiment = {r.sentiment_label: r.count for r in sent_rows.fetchall()}

    # Platform breakdown
    plat_rows = await db.execute(
        select(RawPost.platform, func.count(RawPost.id).label("count"))
        .where(RawPost.collected_at >= since)
        .group_by(RawPost.platform)
    )
    platforms = {r.platform: r.count for r in plat_rows.fetchall()}

    # Critical alerts
    critical_alerts = await db.execute(
        select(Alert).where(Alert.severity.in_(["critical", "high"])).order_by(Alert.created_at.desc()).limit(10)
    )
    critical = [
        {"title": a.title, "severity": a.severity, "type": a.alert_type, "time": a.created_at.strftime("%Y-%m-%d %H:%M")}
        for a in critical_alerts.scalars().all()
    ]

    # Average threat score
    avg_threat = (await db.execute(select(func.avg(ProcessedPost.threat_score)))).scalar() or 0

    return {
        "generated_at": datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"),
        "period": f"Last {days} days",
        "total_posts": total_posts,
        "total_entities": total_entities,
        "alerts_by_severity": alerts_by_severity,
        "top_entities": entities,
        "sentiment": sentiment,
        "platforms": platforms,
        "critical_alerts": critical,
        "avg_threat_score": round(float(avg_threat), 3),
    }


def _sanitize(text: str) -> str:
    """Replace Unicode chars that fpdf can't handle with ASCII equivalents."""
    replacements = {"—": "-", "–": "-", "'": "'", "'": "'", """: '"', """: '"', "…": "...", "•": "*"}
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def generate_pdf(data: dict, filename: str) -> str:
    """Generate E-SitRep as PDF."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Sanitize all string values in data
    def s(text): return _sanitize(str(text))

    # Header
    pdf.set_fill_color(15, 23, 42)  # slate-900
    pdf.rect(0, 0, 210, 40, 'F')
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 24)
    pdf.set_y(8)
    pdf.cell(0, 12, _sanitize("ILA OSINT Intelligence Platform"), ln=True, align="C")
    pdf.set_font("Helvetica", "", 12)
    pdf.cell(0, 8, _sanitize("E-SitRep - Electronic Situation Report"), ln=True, align="C")

    # Classification banner
    pdf.set_fill_color(239, 68, 68)  # red
    pdf.rect(0, 40, 210, 8, 'F')
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_y(41)
    pdf.cell(0, 6, s("CLASSIFICATION: RESTRICTED - FOR AUTHORIZED PERSONNEL ONLY"), ln=True, align="C")

    pdf.set_text_color(0, 0, 0)
    pdf.set_y(55)

    # Report metadata
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(0, 5, f"Generated: {data['generated_at']}  |  Period: {data['period']}", ln=True)
    pdf.ln(5)

    # Executive Summary
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_fill_color(30, 41, 59)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 8, "  1. EXECUTIVE SUMMARY", ln=True, fill=True)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 10)
    pdf.ln(3)

    threat_level = "HIGH" if data['avg_threat_score'] > 0.5 else "MODERATE" if data['avg_threat_score'] > 0.3 else "LOW"
    pdf.multi_cell(0, 5, s(
        f"During the reporting period ({data['period']}), the ILA OSINT platform monitored "
        f"{data['total_posts']:,} posts across {len(data['platforms'])} platforms, tracking "
        f"{data['total_entities']:,} unique entities. The overall threat assessment is {threat_level} "
        f"(score: {data['avg_threat_score']:.1%}). "
        f"{sum(data['alerts_by_severity'].values())} alerts were generated, "
        f"including {data['alerts_by_severity'].get('critical', 0)} critical alerts."
    ))
    pdf.ln(5)

    # Alert Summary
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_fill_color(30, 41, 59)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 8, "  2. ALERT SUMMARY", ln=True, fill=True)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 10)
    pdf.ln(3)

    for sev, count in sorted(data['alerts_by_severity'].items()):
        pdf.cell(60, 6, f"  {sev.upper()}", border=1)
        pdf.cell(30, 6, str(count), border=1, ln=True)
    pdf.ln(5)

    # Critical Alerts
    if data['critical_alerts']:
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_fill_color(30, 41, 59)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(0, 8, "  3. CRITICAL & HIGH PRIORITY ALERTS", ln=True, fill=True)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 9)
        pdf.ln(3)

        for alert in data['critical_alerts']:
            pdf.set_font("Helvetica", "B", 9)
            pdf.cell(20, 5, f"[{alert['severity'].upper()}]")
            pdf.set_font("Helvetica", "", 9)
            pdf.cell(0, 5, s(f"{alert['title'][:80]}"), ln=True)
            pdf.cell(0, 5, f"    Type: {alert['type']}  |  Time: {alert['time']}", ln=True)
            pdf.ln(2)
        pdf.ln(3)

    # Top Threat Entities
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_fill_color(30, 41, 59)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 8, "  4. TOP THREAT ENTITIES", ln=True, fill=True)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 9)
    pdf.ln(3)

    # Table header
    pdf.set_font("Helvetica", "B", 9)
    pdf.cell(25, 6, "Type", border=1)
    pdf.cell(70, 6, "Value", border=1)
    pdf.cell(25, 6, "Risk", border=1)
    pdf.cell(25, 6, "Mentions", border=1, ln=True)
    pdf.set_font("Helvetica", "", 9)

    for entity in data['top_entities']:
        pdf.cell(25, 5, entity['type'], border=1)
        pdf.cell(70, 5, s(str(entity['value'])[:35]), border=1)
        pdf.cell(25, 5, f"{entity['risk']:.0%}", border=1)
        pdf.cell(25, 5, str(entity['mentions']), border=1, ln=True)
    pdf.ln(5)

    # Sentiment Analysis
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_fill_color(30, 41, 59)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 8, "  5. SENTIMENT ANALYSIS", ln=True, fill=True)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 10)
    pdf.ln(3)

    total_sent = sum(data['sentiment'].values()) or 1
    for label, count in data['sentiment'].items():
        pct = count / total_sent * 100
        pdf.cell(40, 6, label.replace('_', ' ').title())
        pdf.cell(30, 6, f"{count:,} ({pct:.0f}%)", ln=True)
    pdf.ln(5)

    # Platform Distribution
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_fill_color(30, 41, 59)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 8, "  6. PLATFORM DISTRIBUTION", ln=True, fill=True)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 10)
    pdf.ln(3)

    for platform, count in sorted(data['platforms'].items(), key=lambda x: -x[1]):
        pdf.cell(40, 6, platform.title())
        pdf.cell(30, 6, f"{count:,}", ln=True)
    pdf.ln(5)

    # Footer
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_fill_color(239, 68, 68)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 8, s("  END OF REPORT - ILA OSINT INTELLIGENCE PLATFORM - RESTRICTED"), ln=True, fill=True, align="C")

    filepath = os.path.join(REPORTS_DIR, filename)
    pdf.output(filepath)
    return filepath


def generate_docx(data: dict, filename: str) -> str:
    """Generate E-SitRep as Word document."""
    doc = Document()

    # Title
    title = doc.add_heading("ILA OSINT Intelligence Platform", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading("E-SitRep — Electronic Situation Report", level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {data['generated_at']}  |  Period: {data['period']}")
    doc.add_paragraph("CLASSIFICATION: RESTRICTED — FOR AUTHORIZED PERSONNEL ONLY").bold = True

    # Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    threat_level = "HIGH" if data['avg_threat_score'] > 0.5 else "MODERATE" if data['avg_threat_score'] > 0.3 else "LOW"
    doc.add_paragraph(
        f"During the reporting period ({data['period']}), the ILA OSINT platform monitored "
        f"{data['total_posts']:,} posts across {len(data['platforms'])} platforms, tracking "
        f"{data['total_entities']:,} unique entities. The overall threat assessment is {threat_level} "
        f"(score: {data['avg_threat_score']:.1%})."
    )

    # Alert Summary
    doc.add_heading("2. Alert Summary", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Severity"
    table.rows[0].cells[1].text = "Count"
    for sev, count in sorted(data['alerts_by_severity'].items()):
        row = table.add_row()
        row.cells[0].text = sev.upper()
        row.cells[1].text = str(count)

    # Critical Alerts
    if data['critical_alerts']:
        doc.add_heading("3. Critical & High Priority Alerts", level=1)
        for alert in data['critical_alerts']:
            doc.add_paragraph(f"[{alert['severity'].upper()}] {alert['title']}", style='List Bullet')

    # Top Entities
    doc.add_heading("4. Top Threat Entities", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    for i, header in enumerate(["Type", "Value", "Risk", "Mentions"]):
        table.rows[0].cells[i].text = header
    for entity in data['top_entities']:
        row = table.add_row()
        row.cells[0].text = entity['type']
        row.cells[1].text = str(entity['value'])[:40]
        row.cells[2].text = f"{entity['risk']:.0%}"
        row.cells[3].text = str(entity['mentions'])

    # Sentiment
    doc.add_heading("5. Sentiment Analysis", level=1)
    total_sent = sum(data['sentiment'].values()) or 1
    for label, count in data['sentiment'].items():
        doc.add_paragraph(f"{label.replace('_', ' ').title()}: {count:,} ({count/total_sent*100:.0f}%)")

    # Platform Distribution
    doc.add_heading("6. Platform Distribution", level=1)
    for platform, count in sorted(data['platforms'].items(), key=lambda x: -x[1]):
        doc.add_paragraph(f"{platform.title()}: {count:,}")

    doc.add_paragraph("\n— END OF REPORT — ILA OSINT INTELLIGENCE PLATFORM — RESTRICTED —")

    filepath = os.path.join(REPORTS_DIR, filename)
    doc.save(filepath)
    return filepath
